####################################################################
# This is a simple Python program for English teaching that uses the AI websites to generate mindmap of given text in English.
# Code on python 3.8.10
# Updated on: 2025-4-17
####################################################################

#IMPORTS
from gooey import Gooey, GooeyParser  # To create a GUI
import argparse  # To parse command line arguments
import os  # To handle file and directory operations
import ctypes  # To interact with Windows API
import json  # To handle JSON data
import webbrowser  # To open web pages
import pyperclip  # To copy and paste text to/from clipboard
import shutil  # To copy and move files and directories
from datetime import datetime  # To get date and time
import xmind  # To create and manipulate XMind files        #limited by xmind.ltd
import platform  # To get the operating system information
import urllib.request  # To handle URL requests
from docx import Document  # To create and manipulate Word documents
from docx.shared import Pt  # To set font size in Word documents
from docx.oxml.ns import qn  # To set font name in Word documents

DOCUMENTS_DIR = os.path.expanduser("~/Documents")
workspace = os.path.join(DOCUMENTS_DIR, "MindDown Docs")
if not os.path.exists(workspace):
    os.makedirs(workspace)
    print(f"Folder Created：{workspace}")
name=""
global arg
arg = None
parse = ''



@Gooey(
        program_name="MindDown 2025 - MindMap Generator",
        program_version="beta 0.1.8", 
        program_description="A simple Python program for English teaching \n" 
        "that uses the AI websites to generate mindmap of given text in English.", 
        default_size=(600, 400), 
        show_restart_button=True,
        language='english'
        ) 
def main():
    models=['ChatGPT-OpenAI','DeepSeek-Official', 'ChatGLM-Tsinghua Edu.','MultiChat-SCNet','Select my own model']

    parser = GooeyParser(description="MindDown 2025 - Mindmap Generator")  # create a parser object
    resize_group = parser.add_argument_group('Options')  # create a group for related options
    resize_group.add_argument(
        '--own_model',
        metavar='Select your own model', 
        help='Select your own model to use for processing',
        type=str,
        default='https://ai.com/'
    )
    resize_group.add_argument(
        '--select_model',
        metavar='Select the model', 
        help='Select the model to use for processing',
        choices=models,
        default=models[0]
    )
    resize_group.add_argument(
        '--word_version',
        metavar='Word version',
        help='Create a word version of the mindmap',
        action='store_true',
        default=True
    )
    global arg  # make arg global to access it later
    arg = parser.parse_args()  # parse the command line arguments

    print("Welcome to MindDown 2025 - Mindmap Generator")
    if arg.select_model == 'Select my own model':
        select_model = arg.own_model
    else:
        select_model = arg.select_model
    print(f"Selected model: {select_model}.\nWord version: {arg.word_version}.")

    config_dir = os.path.join(workspace, "config")
    if not os.path.exists(config_dir):
        os.makedirs(config_dir)
        print(f"Folder Created：{config_dir}")
    pre_file = os.path.join(config_dir,'pre.txt')
    std_file = os.path.join(config_dir, 'std.txt')
    if not os.path.exists(pre_file) or not os.path.exists(std_file):
        print("Something went wrong, please check the config files such as pre.txt and std.txt.")
        raise SystemExit("No such file or directory: pre.txt or std.txt")

    with open(pre_file, 'r', encoding='utf-8') as pre:
        pre = pre.read()
    with open(std_file, 'r', encoding='utf-8') as std:
        std = std.read()
    temp = str(pre) + str(std)
    pyperclip.copy(temp)
    ctypes.windll.user32.MessageBoxW(0, "The pre and std files have been copied to the clipboard. Please paste them into the input box,then upload your text in AI Website.", "MindDown 2025", 1)
    print("The prompt have been copied to the clipboard.")

    if arg.select_model == 'Select my own model':
        webbrowser.open(arg.own_model)
    else:
        if arg.select_model == 'ChatGPT-OpenAI':
            webbrowser.open("https://chatgpt.com/")
        elif arg.select_model == 'DeepSeek-Official':
            webbrowser.open("https://chat.deepseek.com/")
        elif arg.select_model == 'ChatGLM-Tsinghua Edu.':
            webbrowser.open("https://chatglm.cn/")
        elif arg.select_model == 'MultiChat-SCNet':
            webbrowser.open("https://www.scnet.cn/ui/chatbot/")

    current_time = datetime.now().strftime("MindDown_%Y-%m-%d-_%Hh%Mm%Ss")

    if True:
        name = "MindDown_output.xmind"
        output_dir = os.path.join(workspace, current_time, name)
        if os.path.exists(output_dir):
            os.remove(output_dir)
            print(f"File Deleted：{output_dir}")
        if not os.path.exists(os.path.join(workspace, current_time)):
            os.makedirs(os.path.join(workspace, current_time))
            print(f"Folder Created：{os.path.join(workspace, current_time)}")
        parse = json_picker()
        json_2_xmind(parse, output_dir)
        print("Mindmap generated successfully.")

    if arg.word_version:
        name = "MindDown_output.docx"
        output_dir = os.path.join(workspace, current_time, name)
        if not os.path.exists(os.path.join(workspace, current_time)):
            os.makedirs(os.path.join(workspace, current_time))
            print(f"Folder Created：{os.path.join(workspace, current_time)}")
        doc = json_to_docx(parse, font_name="宋体", font_size=9)
        doc.save(output_dir)
        print("Word version generated successfully.")

    return "end"

'''FUCTIONS'''
# Check if the URL is reachable
def check_connectivity(url,t=5): #code by Lingxian Wrang
    try:
        response = urllib.request.urlopen(url, timeout=t)
        return True
    except Exception as e:
        return False
    
def json_picker(): #code by Lingxian Wrang
    while True:
        response = ctypes.windll.user32.MessageBoxW(0, "Please copy the JSON given by AI and press OK.", "MindDown 2025", 1)
        if response == 2:  # If the user clicks "Cancel"
            print("Operation cancelled by user: JSON input cancelled.")
            raise SystemExit("Operation cancelled by user: JSON input cancelled.")
        clipboard_content = pyperclip.paste()  # Get the content from the clipboard
        if not clipboard_content:  # If the clipboard is empty
            ctypes.windll.user32.MessageBoxW(0, "Clipboard seems to be empty.Plz try again.", "MindDown 2025", 0x10)
            print("Error: Clipboard is empty.")
        try:
            json_data = json.loads(clipboard_content)
            parse_file = os.path.join(workspace,'%temp%','parse.json')
            if not os.path.exists(os.path.join(workspace,'%temp%')):
                os.makedirs(os.path.join(workspace,'%temp%'))
                print(f"Folder Created：{os.path.join(workspace,'%temp%')}")
            with open(parse_file, 'w', encoding='utf-8') as file:
                json.dump(json_data, file, ensure_ascii=False, indent=4)
            print("Success: parsed the JSON data.")
            break
        except json.JSONDecodeError:  # If the JSON is not valid
            ctypes.windll.user32.MessageBoxW(0, "Something went wrong. Plz check JSON then copy it and try again.", "MindDown2025", 0x10)
            print("Error: JSON data is not valid.")
    with open(parse_file, 'r', encoding='utf-8') as file:
        temp = file.read()
        return temp

def genXmindByJson(parent, data): #code by Github user zhaojinzhou
    theTree = data
    if theTree is None:
        return
    node = parent.addSubTopic()
    for key in theTree:
        print(key)
        if key == 'name':
            print(theTree['name'])
            node.setTitle(theTree['name'])
        elif key == 'children':
            if len(theTree['children'])>0:
                for son_tree in theTree['children']:
                    genXmindByJson(node,son_tree)
            else:
                pass
        else:
            pass
    print("Xmind_SDK: Done.")

def genXmind(input): #code by Github user zhaojinzhou
    workbook = xmind.load("temp.xmind")
    sheet = workbook.getPrimarySheet()
    root = sheet.getRootTopic()
    root.setTitle(input['name'])
    if 'children' in input:
        for son_tree in input['children']:
            genXmindByJson(root, son_tree)
    xmind.save(workbook=workbook, path=input['name'] + '.xmind')

def dfs(topic_data, topic): #code by Github user zhaojinzhou
    if isinstance(topic_data,dict):
        for i in topic_data.keys():
            sub_topic = topic.addSubTopic()
            sub_topic.setTitle(i)
            dfs(topic_data[i], sub_topic)
    else:
        if isinstance(topic_data,list):
            data=''
            index = 0
            for i in topic_data:
                if isinstance(i,dict):
                    sub_topic = topic.addSubTopic()
                    sub_topic.setTitle(index)
                    index += 1
                    dfs(i, sub_topic)
                    continue
                data += i
                data += ','
            if(len(data) != 0):
                sub_topic = topic.addSubTopic()
                sub_topic.setTitle(data)
        else:
            sub_topic = topic.addSubTopic()
            sub_topic.setTitle(topic_data)
            return sub_topic

def json_2_xmind(data, xmind_name): #code by Hira
    workbook = xmind.load(xmind_name)
    sheet1 = workbook.getPrimarySheet()
    sheet1.setTitle(xmind_name)
    root_topic1 = sheet1.getRootTopic()
    root_topic1.setTitle(xmind_name)
    dfs(data, root_topic1)
    xmind.save(workbook, path=xmind_name)

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def add_paragraph_with_font(doc, text, font_name="宋体", font_size=12):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # 设置中文字体（尤其是 eastAsia）
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return paragraph

def json_to_docx(data, doc=None, indent=0, font_name="宋体", font_size=9):
    if doc is None:
        doc = Document()

        style = doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    indent_str = "    " * indent

    if isinstance(data, dict):
        for key, value in data.items():
            add_paragraph_with_font(doc, f"{indent_str}{key}:", font_name, font_size)
            json_to_docx(value, doc, indent + 1, font_name, font_size)
    elif isinstance(data, list):
        for index, item in enumerate(data):
            add_paragraph_with_font(doc, f"{indent_str}- 项 {index + 1}：", font_name, font_size)
            json_to_docx(item, doc, indent + 1, font_name, font_size)
    else:
        add_paragraph_with_font(doc, f"{indent_str}{data}", font_name, font_size)

    return doc


if __name__ == "__main__":
    main()
